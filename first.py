from docx2pdf import convert
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from flask import Flask, request, send_file, jsonify
import os
import pandas as pd
from openai import OpenAI
import tempfile
import json
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

client = OpenAI(api_key="AHTC6qXxqJJUCy8u9TnStLzrrjOsyqMT", base_url="https://api.deepinfra.com/v1/openai")


def generate_response_from_api(system_prompt: str, prompt: str, temp=0.8):
    messages = [
        {'role': 'system', 'content': system_prompt},
        {'role': 'user', 'content': prompt}
    ]

    response = client.chat.completions.create(
        model="meta-llama/Meta-Llama-3-70B-Instruct",
        messages=messages,
        temperature=temp,
        max_tokens=300
    )

    if response.choices:
        return response.choices[0].message.content
    return ""

def replace_text(doc, old_text, new_text):
    found = False
    for paragraph in doc.paragraphs:
        if found:
            break
        if old_text in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if old_text in item.text:
                    item.text = item.text.replace(old_text, new_text, 1)
                    found = True
                    break

    for table in doc.tables:
        if found:
            break
        for row in table.rows:
            if found:
                break
            for cell in row.cells:
                if found:
                    break
                if old_text in cell.text:
                    replace_text_in_cell(cell, old_text, new_text)
                    found = True

def replace_text_in_cell(cell, old_text, new_text):
    for paragraph in cell.paragraphs:
        if old_text in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if old_text in item.text:
                    item.text = item.text.replace(old_text, new_text, 1)
                    return

app = Flask(__name__)




@app.route('/cv', methods=['POST'])
def convert_vector():
    data = request.get_json()

    name = data['name']
    age = data['age']
    description = data['description']
    education = data['education']
    location = data['location']
    country = data['country']
    move = data['move']
    business_trip = data['businessTrip']
    schedule = data['sheduel']
    prev_job = data['prevJob']
    skills = data['skills']

    system_prompt = "Select from list which IT profession the person belongs to [.net, AT, BA, DevOps, Frontend, Functional Testing, Java, Python, PM]. Select only one and write it with level of person (junior, mid, senior) evaluate a person strictly. Respond with only two words: the profession and the level."
    prompt = description + ' ' + education + ' ' + prev_job

    job = generate_response_from_api(system_prompt, prompt, temp=0.3)
    print(job)

    system_prompt = "Ты человек, который пишет резюме. Напишите короткий текст из 8-13 слов о своем опыте. Ответ должен быть кратким и основанным на фактах, без дополнительной информации или личной интерпретации. Пишите с точки зрения пользователя и заключайте ответ в квадратные скобки."
    prompt = description + ' ' + education + ' ' + prev_job + ' ' + skills

    tag1 = generate_response_from_api(system_prompt, prompt, temp=0.3)
    tag1 = re.findall(r'\[(.*?)\]', tag1)
    tag1 = tag1[0]
    print(tag1)

    system_prompt = "Ты человек, который пишет резюме. Напишите короткий текст из 8-13 слов о среднем количество лет опыта, домены и проекты, в которых ты работал. Ответ должен быть кратким и основанным на фактах, без дополнительной информации или личной интерпретации, если данной информации нет, не пиши ничего. Пишите с точки зрения пользователя и заключайте ответ в квадратные скобки."
    prompt = description + ' ' + education + ' ' + prev_job + ' ' + skills

    tag2 = generate_response_from_api(system_prompt, prompt, temp=0.3)
    tag2 = re.findall(r'\[(.*?)\]', tag2)
    tag2 = tag2[0]
    print(tag2)

    system_prompt = "Напиши короткий текст из 5-8 слов о твоих сильны навыках и сторонах. Ответ должен быть кратким и основанным на фактах, без дополнительной информации или личной интерпретации, если данной информации нет, не пиши ничего. Пишите с точки зрения пользователя и заключайте ответ в квадратные скобки."
    prompt = description + ' ' + education + ' ' + prev_job + ' ' + skills

    tag3 = generate_response_from_api(system_prompt, prompt, temp=0.3)
    tag3 = re.findall(r'\[(.*?)\]', tag3)
    tag3 = tag3[0]
    print(tag3)

    system_prompt = "Напиши текст из 15-20 слов о твоем списке основных инструментов, технологий. Ответ должен быть основанным на фактах, без дополнительной информации или личной интерпретации, если данной информации нет, не пиши ничего. Пишите с точки зрения пользователя и заключайте ответ в квадратные скобки."
    prompt = description + ' ' + education + ' ' + prev_job + ' ' + skills

    tag4 = generate_response_from_api(system_prompt, prompt, temp=0.3)
    tag4 = re.findall(r'\[(.*?)\]', tag4)
    tag4 = tag4[0]
    print(tag4)

    system_prompt = "Выбери высшее или среднее образование. Ответ должен быть основанным на фактах, без дополнительной информации или личной интерпретации, если данной информации нет, не пиши ничего. Ответ должен быть 1 словом (высшее или среднее) ответ в квадратные скобки."
    prompt = education

    tag5 = generate_response_from_api(system_prompt, prompt, temp=0.3)
    tag5 = re.findall(r'\[(.*?)\]', tag5)
    tag5 = tag5[0]
    print(tag5)

    system_prompt = "Напиши специальность в 2-3 словах. Ответ должен быть основанным на фактах, без дополнительной информации или личной интерпретации, если данной информации нет, не пиши ничего. Пиши ответ в квадратных скобках."
    prompt = education

    tag6 = generate_response_from_api(system_prompt, prompt, temp=0.3)
    tag6 = re.findall(r'\[(.*?)\]', tag6)
    tag6 = tag6[0]
    print(tag6)

    system_prompt = "Напиши название учебного заведения. Ответ должен быть основанным на фактах, без дополнительной информации или личной интерпретации, если данной информации нет, не пиши ничего. Пиши ответ в квадратных скобках."
    prompt = education

    tag7 = generate_response_from_api(system_prompt, prompt, temp=0.3)
    tag7 = re.findall(r'\[(.*?)\]', tag7)
    tag7 = tag7[0]
    print(tag7)

    doc_path = 'CV_template.docx'
    doc = Document(doc_path)

    replace_text(doc, 'Имя', name)
    replace_text(doc, 'Должность', job)
    replace_text(doc, 'Обобщение опыта', tag1)
    replace_text(doc, 'Среднее количество лет опыта, домены и проекты, в которых работал кандидат', tag2)
    replace_text(doc, 'Сильные стороны и хорошо развитые навыки', tag3)

    replace_text(doc, 'Перечисление основных инструментов, технологий, с которыми работал кандидат', tag4)

    replace_text(doc, 'Высшее/среднее образование', tag5 + " образование")

    replace_text(doc, 'Специальность', tag6)
    replace_text(doc, 'Учебное заведение', tag7)

    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_docx.name)
    temp_docx.close()

    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    temp_pdf.close()

    # Конвертируем DOCX в PDF
    convert(temp_docx.name, temp_pdf.name)

    with open(temp_pdf.name, 'rb') as pdf_file:
        pdf_bytes = pdf_file.read()

    os.remove(temp_pdf.name)

    return pdf_bytes  # Возвращаем байтовый массив PDF файла

if __name__ == '__main__':
    app.run(debug=True, port=3000)